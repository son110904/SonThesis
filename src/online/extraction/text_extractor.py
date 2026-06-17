"""
text_extractor.py – Trích xuất văn bản thô từ CV (PDF hoặc DOCX).

Bước 2 của Online Pipeline.

Output: raw_text (str)
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx"}


class UnsupportedFileType(ValueError):
    """Raise khi file không phải PDF/DOCX."""


def _extract_pdf(data: bytes) -> str:
    """Trích text từ PDF bytes bằng pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001 - 1 page lỗi không nên hỏng cả file
            logger.warning(f"Lỗi trích 1 trang PDF: {e}")
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    """Trích text từ DOCX bytes bằng python-docx (cả paragraph + bảng)."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

    # Lấy thêm text trong bảng (CV hay dùng bảng để layout)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """
    Trích văn bản thô từ nội dung file (bytes).

    Args:
        data:     Nội dung file dạng bytes.
        filename: Tên file (để xác định loại qua phần mở rộng).

    Returns:
        Văn bản thô đã trích, đã strip.

    Raises:
        UnsupportedFileType: Nếu không phải PDF/DOCX.
    """
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"Định dạng '{ext}' không hỗ trợ. Chỉ chấp nhận: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    if ext == ".pdf":
        text = _extract_pdf(data)
    else:
        text = _extract_docx(data)

    text = text.strip()
    logger.info(f"Trích xuất '{filename}' ({ext}) → {len(text)} ký tự")
    if not text:
        logger.warning(f"File '{filename}' trích ra rỗng — có thể là PDF scan ảnh.")
    return text


def extract_text(file_path: str | Path) -> str:
    """
    Trích văn bản thô từ đường dẫn file trên đĩa.

    Args:
        file_path: Đường dẫn tới file PDF/DOCX.

    Returns:
        Văn bản thô đã trích.

    Raises:
        FileNotFoundError:   Nếu file không tồn tại.
        UnsupportedFileType: Nếu không phải PDF/DOCX.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Không tìm thấy file: {path}")
    data = path.read_bytes()
    return extract_text_from_bytes(data, path.name)
