"""
text_extractor.py – Trích xuất văn bản thô từ CV (PDF hoặc DOCX).

Bước 2 của Online Pipeline.

Output: raw_text (str)
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".md"}

# ── Column-aware PDF extraction (xem _extract_pdf) ───────────────────────
# CV 2 cột (sidebar Contact/Skill bên trái + nội dung chính bên phải) khiến
# pypdf trả text sai thứ tự (theo thứ tự object trong content-stream, không
# phải vị trí hiển thị trên trang). Các hằng số dưới đây điều khiển việc dò
# cột bằng khoảng trắng dọc (gutter) giữa 2 cột, dựa trên bounding box của
# từng dòng — dựng trực tiếp từ vị trí (x, y) của từng ký tự (pdfplumber
# `page.chars`), không dùng word-tokenizer của pdfplumber vì tolerance cố
# định của nó dễ vỡ với PDF có tracking chữ bất thường (Canva/Figma) hoặc
# PDF không có khoảng trắng thật giữa các từ (LaTeX/Overleaf).
_LINE_TOP_TOL = 3.0        # dung sai (pt) để gộp ký tự cùng hàng ngang thành 1 "hàng"
_GAP_SIZE_RATIO = 0.15     # khoảng cách > tỉ lệ này * cỡ chữ → coi là ranh giới từ (thêm space)
_LINE_SPLIT_SIZE_RATIO = 2.0  # khoảng cách > tỉ lệ này * cỡ chữ → tách thành 2 dòng riêng biệt
_SPAN_WIDTH_FRAC = 0.6     # dòng rộng hơn 60% trang → coi là tiêu đề/spanning, không thuộc cột nào
_MIN_GUTTER_WIDTH = 12.0   # bề rộng tối thiểu (pt) của khoảng trắng giữa 2 cột
_MARGIN_FRAC = 0.12        # bỏ qua khoảng trắng quá gần lề trái/phải (không phải gutter thật)
_MIN_LINES_PER_SIDE = 3    # cần tối thiểu N dòng mỗi bên mới tin là layout 2 cột
_SAME_ROW_ARTIFACT_FRAC = 0.4  # ≥ tỉ lệ này dòng bên nhỏ trùng top với bên kia → coi là badge/label cùng hàng, không phải 2 cột thật


def _looks_char_spaced(line: str) -> bool:
    """
    True nếu dòng bị "giãn ký tự" — mỗi chữ cái cách nhau 1 space
    (vd 'P y t h o n'), thường gặp ở PDF xuất từ Canva/Figma.

    Heuristic: ≥45% token (tách bởi 1 space) là 1 ký tự CHỮ/SỐ. Loại ký tự
    đơn không phải chữ/số (icon Font Awesome, dấu '|', bullet…) khỏi phép
    đếm — dòng contact info kiểu " Espoo, Finland  email@x.com" (icon
    xen giữa các mục) không phải char-spacing thật, không nên bị vá.
    """
    tokens = [t for t in line.split(" ") if t != ""]
    if len(tokens) < 8:
        return False
    single = sum(1 for t in tokens if len(t) == 1 and t.isalnum())
    return single / len(tokens) >= 0.45


def _repair_char_spacing(text: str) -> str:
    """
    Sửa text bị giãn ký tự. Quy luật của loại PDF này:
        - trong 1 TỪ: các ký tự cách nhau 1 space  → 'P y t h o n'
        - giữa 2 TỪ : cách nhau 2+ space           → 'P y t h o n   F a s t A P I'
    Nên: tách dòng theo run 2+ space (ranh giới từ), rồi bỏ space đơn bên trong.
    Chỉ áp dụng cho dòng thực sự bị giãn, dòng bình thường giữ nguyên.
    """
    out: list[str] = []
    fixed_any = False
    for line in text.split("\n"):
        if _looks_char_spaced(line):
            words = re.split(r" {2,}", line.strip())
            words = [w.replace(" ", "") for w in words]
            out.append(" ".join(w for w in words if w))
            fixed_any = True
        else:
            out.append(line)
    if fixed_any:
        logger.info("Phát hiện & sửa PDF bị giãn ký tự (char-spaced).")
    return "\n".join(out)


class UnsupportedFileType(ValueError):
    """Raise khi file không phải PDF/DOCX."""


def _join_chars(cs: list[dict[str, Any]]) -> str:
    """
    Ghép 1 dãy ký tự (đã sort theo x0, cùng 1 "run") thành text, chèn space
    khi: ký tự đó tự thân là whitespace (space thật trong PDF), HOẶC
    khoảng cách giữa 2 ký tự vượt _GAP_SIZE_RATIO * cỡ chữ (bắt các từ được
    đặt vị trí bằng offset thay vì ký tự space thật, vd PDF từ LaTeX/Overleaf).
    """
    parts: list[str] = []
    prev_x1: float | None = None
    prev_size: float | None = None
    prev_was_space = True
    for c in cs:
        txt = c["text"]
        if txt.isspace():
            if not prev_was_space:
                parts.append(" ")
            prev_was_space = True
        else:
            if prev_x1 is not None and not prev_was_space:
                gap = c["x0"] - prev_x1
                avg_size = ((prev_size or c["size"]) + c["size"]) / 2
                if gap > _GAP_SIZE_RATIO * avg_size:
                    parts.append(" ")
            parts.append(txt)
            prev_was_space = False
        prev_x1 = c["x1"]
        prev_size = c["size"]
    return "".join(parts).strip()


def _chars_to_lines(chars: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """
    Gộp các ký tự (pdfplumber `page.chars`) thành các dòng, mỗi dòng có
    text/x0/x1/top.

    2 bước:
    1. Gộp ký tự cùng hàng ngang (top gần nhau) thành 1 "hàng" thô.
    2. Trong mỗi hàng thô, tách thành nhiều dòng riêng nếu có khoảng cách
       ngang lớn (> _LINE_SPLIT_SIZE_RATIO * cỡ chữ) giữa 2 ký tự liền kề —
       trường hợp này thường là 2 cột khác nhau vô tình có cùng độ cao
       (vd tiêu đề "CONTACT" bên trái và "PROFESSIONAL SUMMARY" bên phải
       nằm cùng 1 hàng), hoặc 1 badge ngày tháng canh phải trên cùng dòng
       với nội dung bên trái — cả 2 trường hợp đều KHÔNG nên bị nối chung.
    """
    rows: list[dict[str, Any]] = []
    for c in sorted(chars, key=lambda c: (c["top"], c["x0"])):
        row = next((r for r in rows if abs(r["top"] - c["top"]) <= _LINE_TOP_TOL), None)
        if row is None:
            row = {"top": c["top"], "chars": []}
            rows.append(row)
        row["chars"].append(c)

    out: list[dict[str, Any]] = []
    for row in rows:
        cs = sorted(row["chars"], key=lambda c: c["x0"])
        run: list[dict[str, Any]] = []
        prev_x1: float | None = None
        prev_size: float | None = None
        for c in cs:
            if prev_x1 is not None and not c["text"].isspace():
                gap = c["x0"] - prev_x1
                avg_size = ((prev_size or c["size"]) + c["size"]) / 2
                if gap > _LINE_SPLIT_SIZE_RATIO * avg_size and run:
                    text = _join_chars(run)
                    if text:
                        out.append({
                            "top": min(c["top"] for c in run),
                            "x0": min(c["x0"] for c in run),
                            "x1": max(c["x1"] for c in run),
                            "text": text,
                        })
                    run = []
            run.append(c)
            prev_x1 = c["x1"]
            prev_size = c["size"]
        if run:
            text = _join_chars(run)
            if text:
                out.append({
                    "top": min(c["top"] for c in run),
                    "x0": min(c["x0"] for c in run),
                    "x1": max(c["x1"] for c in run),
                    "text": text,
                })
    out.sort(key=lambda l: (l["top"], l["x0"]))
    return out


def _partition_columns(
    lines: list[dict[str, Any]], gutter_x: float, page_width: float
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    """Chia các dòng thành (trái, phải, spanning) theo tọa độ gutter_x."""
    span_width = _SPAN_WIDTH_FRAC * page_width
    left: list[dict[str, Any]] = []
    right: list[dict[str, Any]] = []
    spanning: list[dict[str, Any]] = []
    for l in lines:
        if (l["x1"] - l["x0"]) >= span_width:
            spanning.append(l)
            continue
        left_overlap = max(0.0, min(l["x1"], gutter_x) - l["x0"])
        right_overlap = max(0.0, l["x1"] - max(l["x0"], gutter_x))
        (left if left_overlap >= right_overlap else right).append(l)
    return left, right, spanning


def _is_same_row_artifact(left: list[dict[str, Any]], right: list[dict[str, Any]]) -> bool:
    """
    True nếu phần lớn dòng ở 1 bên thực ra chỉ là mảnh vỡ tách ra từ CÙNG 1
    hàng ngang với 1 dòng ở bên kia (top gần như trùng khớp) — dấu hiệu của
    badge ngày tháng canh phải trên 1 CV thật ra chỉ có 1 cột (vd "Software
    Engineer Intern .......... June 2024 – Aug 2024"), KHÔNG phải layout 2
    cột thật (2 cột thật: nội dung mỗi bên trôi độc lập, top hiếm khi trùng).
    """
    smaller, other = (left, right) if len(left) <= len(right) else (right, left)
    if not smaller:
        return False
    paired = sum(
        1 for l in smaller if any(abs(l["top"] - o["top"]) <= _LINE_TOP_TOL for o in other)
    )
    return (paired / len(smaller)) >= _SAME_ROW_ARTIFACT_FRAC


def _find_gutter(lines: list[dict[str, Any]], page_width: float) -> float | None:
    """
    Tìm khoảng trắng dọc (gutter) ngăn cách 2 cột, dựa trên các dòng "hẹp"
    (không phải tiêu đề/spanning). Trả về tọa độ x giữa gutter, hoặc None
    nếu trang không đủ bằng chứng là layout 2 cột (an toàn: coi là 1 cột).
    """
    narrow = [l for l in lines if (l["x1"] - l["x0"]) < _SPAN_WIDTH_FRAC * page_width]
    if not narrow:
        return None

    intervals = sorted((l["x0"], l["x1"]) for l in narrow)
    merged: list[list[float]] = []
    for x0, x1 in intervals:
        if merged and x0 <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], x1)
        else:
            merged.append([x0, x1])

    lo, hi = page_width * _MARGIN_FRAC, page_width * (1 - _MARGIN_FRAC)
    gaps = [
        (merged[i][1], merged[i + 1][0])
        for i in range(len(merged) - 1)
        if merged[i + 1][0] - merged[i][1] >= _MIN_GUTTER_WIDTH
        and merged[i][1] >= lo
        and merged[i + 1][0] <= hi
    ]
    if not gaps:
        return None

    gstart, gend = max(gaps, key=lambda g: g[1] - g[0])
    gutter_x = (gstart + gend) / 2
    left, right, _ = _partition_columns(narrow, gutter_x, page_width)
    if len(left) < _MIN_LINES_PER_SIDE or len(right) < _MIN_LINES_PER_SIDE:
        return None
    if _is_same_row_artifact(left, right):
        return None
    return gutter_x


def _order_lines(lines: list[dict[str, Any]], page_width: float) -> list[str]:
    """
    Sắp lại thứ tự đọc của các dòng trong 1 trang.

    Nếu phát hiện layout 2 cột: đọc hết cột trái (sidebar) rồi đến cột phải
    (nội dung chính), theo thứ tự trên-xuống trong mỗi cột. Dòng "spanning"
    (span >= 60% bề rộng trang, ví dụ tên/tiêu đề) xuất hiện trước cột (nếu
    ở đầu trang) hoặc sau cùng (nếu nằm giữa/cuối trang).

    Nếu không đủ bằng chứng 2 cột: giữ nguyên thứ tự trên-xuống mặc định
    (pdfplumber đã sắp theo `top`).
    """
    if not lines:
        return []

    gutter_x = _find_gutter(lines, page_width)
    if gutter_x is None:
        return [l["text"] for l in lines]

    left, right, spanning = _partition_columns(lines, gutter_x, page_width)

    first_col_top = min(l["top"] for l in (left + right))
    preamble = [l for l in spanning if l["top"] < first_col_top]
    trailing = [l for l in spanning if l["top"] >= first_col_top]

    ordered = preamble + left + right + trailing
    return [l["text"] for l in ordered]


def _extract_pdf_pypdf(data: bytes) -> str:
    """Trích text từ PDF bytes bằng pypdf (fallback khi pdfplumber lỗi)."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001 - 1 page lỗi không nên hỏng cả file
            logger.warning(f"Lỗi trích 1 trang PDF: {e}")
    return "\n".join(pages)


def _extract_pdf(data: bytes) -> str:
    """
    Trích text từ PDF bytes, có nhận diện layout 2 cột (sidebar CV).

    Dùng pdfplumber để lấy vị trí (bounding box) từng ký tự, dựng lại dòng
    (xem _chars_to_lines), dò cột và sắp lại đúng thứ tự đọc (_order_lines).
    Nếu pdfplumber lỗi (PDF hỏng, không đọc được…) → rơi về pypdf như cũ.
    """
    try:
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                try:
                    lines = _chars_to_lines(page.chars)
                    pages.append("\n".join(_order_lines(lines, page.width)))
                except Exception as e:  # noqa: BLE001 - 1 trang lỗi không nên hỏng cả file
                    logger.warning(f"Lỗi trích 1 trang PDF (pdfplumber): {e}")
        return "\n".join(pages)
    except Exception as e:  # noqa: BLE001 - pdfplumber lỗi toàn bộ → fallback pypdf
        logger.warning(f"pdfplumber lỗi, fallback sang pypdf: {e}")
        return _extract_pdf_pypdf(data)


def _extract_md(data: bytes) -> str:
    """
    Trích text từ Markdown bytes. Bỏ cú pháp MD cơ bản để regex skill chạy sạch
    (heading #, bullet -/*, bold **, code `, link [text](url) → text).
    """
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="ignore")

    # [text](url) → text  ; ![alt](url) → alt
    text = re.sub(r"!?\[([^\]]*)\]\([^)]*\)", r"\1", text)
    # bỏ ký tự cú pháp đầu/giữa dòng, giữ nội dung
    text = re.sub(r"^[ \t]*#{1,6}\s*", "", text, flags=re.MULTILINE)  # heading
    text = re.sub(r"^[ \t]*[-*+]\s+", "", text, flags=re.MULTILINE)   # bullet
    text = re.sub(r"^[ \t]*>\s?", "", text, flags=re.MULTILINE)        # blockquote
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text


def _extract_docx(data: bytes) -> str:
    """Trích text từ DOCX bytes bằng python-docx (cả paragraph + bảng)."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

    # Lấy thêm text trong bảng (CV hay dùng bảng để layout)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """
    Trích văn bản thô từ nội dung file (bytes).

    Args:
        data:     Nội dung file dạng bytes.
        filename: Tên file (để xác định loại qua phần mở rộng).

    Returns:
        Văn bản thô đã trích, đã strip.

    Raises:
        UnsupportedFileType: Nếu không phải PDF/DOCX.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".doc":
        # .doc (Word nhị phân cũ) không đọc được bằng Python thuần → hướng dẫn chuyển đổi.
        raise UnsupportedFileType(
            "Định dạng .doc (Word cũ) không được hỗ trợ. "
            "Vui lòng lưu lại dưới dạng .docx, PDF hoặc .md rồi tải lên."
        )
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"Định dạng '{ext}' không hỗ trợ. Chỉ chấp nhận: PDF, DOCX, Markdown (.md)."
        )

    if ext == ".pdf":
        text = _extract_pdf(data)
        text = _repair_char_spacing(text)  # vá PDF giãn ký tự (Canva/Figma…)
    elif ext == ".md":
        text = _extract_md(data)
    else:
        text = _extract_docx(data)

    text = text.strip()
    logger.info(f"Trích xuất '{filename}' ({ext}) → {len(text)} ký tự")
    if not text:
        logger.warning(f"File '{filename}' trích ra rỗng — có thể là PDF scan ảnh.")
    return text


def extract_text(file_path: str | Path) -> str:
    """
    Trích văn bản thô từ đường dẫn file trên đĩa.

    Args:
        file_path: Đường dẫn tới file PDF/DOCX.

    Returns:
        Văn bản thô đã trích.

    Raises:
        FileNotFoundError:   Nếu file không tồn tại.
        UnsupportedFileType: Nếu không phải PDF/DOCX.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Không tìm thấy file: {path}")
    data = path.read_bytes()
    return extract_text_from_bytes(data, path.name)
